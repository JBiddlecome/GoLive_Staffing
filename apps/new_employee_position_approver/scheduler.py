from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import mimetypes
import os
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

import docx
import httpx
import openai
import pypdf
from sqlalchemy import text

from apps.position_requests.scheduler import LEVEL_MAP, POSITION_KEY_TO_NAME, _engine
from apps.resume_analyzer.views import RESUME_SYSTEM_PROMPT

logger = logging.getLogger(__name__)

DATA_FILE = Path("data/new_employee_position_approver.json")
CONFIG_FILE = Path("data/new_employee_position_approver_config.json")
S3_BUCKET = os.getenv("S3_BUCKET", "web-application-files")
S3_REGION = os.getenv("S3_REGION", "us-east-1")
S3_PREFIX = os.getenv("RESUME_S3_PREFIX", "employee/resume/")

_auto_approve_enabled = False
_state_initialized = False

def _ensure_state_table():
    from sqlalchemy import text
    from apps.position_requests.scheduler import _engine
    engine = _engine()
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS golive_app_state (
                    app_name VARCHAR(100) NOT NULL,
                    state_key VARCHAR(100) NOT NULL,
                    state_value LONGTEXT,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    PRIMARY KEY (app_name, state_key)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """))
    except Exception as e:
        logger.error(f"[New Employee Position Approver] Error ensuring state table: {e}")
    finally:
        engine.dispose()

def get_automation_enabled() -> bool:
    global _auto_approve_enabled, _state_initialized
    if not _state_initialized:
        _ensure_state_table()
        from sqlalchemy import text
        from apps.position_requests.scheduler import _engine
        engine = _engine()
        try:
            with engine.begin() as conn:
                res = conn.execute(text("SELECT state_value FROM golive_app_state WHERE app_name = 'new_employee_position_approver' AND state_key = 'auto_approve'")).fetchone()
                if res and res[0] is not None:
                    _auto_approve_enabled = (res[0].lower() == 'true')
        except Exception as e:
            logger.error(f"[New Employee Position Approver] Error loading auto approve state from DB: {e}")
        finally:
            engine.dispose()
        _state_initialized = True
    return _auto_approve_enabled

def set_automation_enabled(enabled: bool) -> None:
    global _auto_approve_enabled, _state_initialized
    _auto_approve_enabled = enabled
    _state_initialized = True
    _ensure_state_table()
    from sqlalchemy import text
    from apps.position_requests.scheduler import _engine
    engine = _engine()
    try:
        with engine.begin() as conn:
            val_str = 'true' if enabled else 'false'
            conn.execute(text("""
                INSERT INTO golive_app_state (app_name, state_key, state_value)
                VALUES ('new_employee_position_approver', 'auto_approve', :value)
                ON DUPLICATE KEY UPDATE state_value = :value
            """), {"value": val_str})
    except Exception as e:
        logger.error(f"[New Employee Position Approver] Error saving auto approve state to DB: {e}")
    finally:
        engine.dispose()

_state_lock = asyncio.Lock()

LEVELED_POSITION_KEYS = {"cook", "prep_cook", "server", "bartender", "barista"}
MAX_AUTO_APPROVAL_LEVEL = 2

# Only these 18 Resume Analyzer position families are eligible for this app.
# Values are the exact active database position descriptions to add.
APPROVAL_POSITION_DESCRIPTIONS: dict[str, str | dict[int, str]] = {
    "cook": {1: "Cook 1", 2: "Cook 2"},
    "prep_cook": {1: "Prep Cook 1", 2: "Prep Cook 2"},
    "dishwasher": "Dishwasher",
    "utility": "Utility",
    "server": {1: "Server 1", 2: "Server 2"},
    "host": "Host",
    "runner": "Runner",
    "busser": "Busser",
    "bartender": {1: "Bartender 1", 2: "Bartender 2"},
    "barback": "Barback",
    "cashier": "Cashier",
    "pastry": "Pastry Cook",
    "baker": "Baker",
    "sushi": "Sushi Chef 1",
    "concessions": "Concessions Worker",
    "barista": {1: "Barista 1", 2: "Barista 2"},
    "valet": "Valet",
    "event_supervisor": "Event Supervisor",
    "sous_chef_2": "Sous Chef 2",
}


def _now_iso() -> str:
    return _local_now().replace(microsecond=0).isoformat()


def _today() -> str:
    return _local_now().date().isoformat()


def _local_now() -> datetime:
    timezone_name = os.getenv("APP_TIMEZONE", "America/Los_Angeles")
    try:
        return datetime.now(ZoneInfo(timezone_name))
    except ZoneInfoNotFoundError:
        return datetime.now()


def _empty_state() -> dict[str, Any]:
    return {"reviewed": {}, "records": [], "last_scan_at": None, "last_scan_message": ""}


def load_state() -> dict[str, Any]:
    if not DATA_FILE.exists():
        return _empty_state()

    try:
        state = json.loads(DATA_FILE.read_text(encoding="utf-8"))
    except Exception:
        logger.exception("Could not read New Employee Position Approver state")
        return _empty_state()

    if not isinstance(state, dict):
        return _empty_state()

    state.setdefault("reviewed", {})
    state.setdefault("records", [])
    state.setdefault("last_scan_at", None)
    state.setdefault("last_scan_message", "")
    return state


def save_state(state: dict[str, Any]) -> None:
    DATA_FILE.parent.mkdir(parents=True, exist_ok=True)
    payload = json.dumps(state, indent=2, default=str)
    temp_file = DATA_FILE.with_name(f"{DATA_FILE.name}.tmp")
    temp_file.write_text(payload, encoding="utf-8")
    temp_file.replace(DATA_FILE)


def _row_value(row: Any, key: str, index: int | None = None) -> Any:
    if hasattr(row, "_mapping") and key in row._mapping:
        return row._mapping[key]
    if hasattr(row, key):
        return getattr(row, key)
    if index is not None:
        return row[index]
    return None


def _resume_url(filename: str) -> str:
    safe_filename = quote(filename)
    return f"https://{S3_BUCKET}.s3.{S3_REGION}.amazonaws.com/{S3_PREFIX}{safe_filename}"


def _allowed_database_descriptions() -> set[str]:
    descriptions: set[str] = set()
    for value in APPROVAL_POSITION_DESCRIPTIONS.values():
        if isinstance(value, dict):
            descriptions.update(value.values())
        else:
            descriptions.add(value)
    return descriptions


def _get_allowed_catalog(conn) -> dict[str, dict[str, Any]]:
    allowed = _allowed_database_descriptions()
    rows = conn.execute(
        text(
            """
            SELECT position_id, TRIM(description) AS description
            FROM position
            WHERE deleted_at IS NULL
              AND description IS NOT NULL
              AND description != ''
            ORDER BY position_id
            """
        )
    ).mappings().all()

    catalog: dict[str, dict[str, Any]] = {}
    for row in rows:
        description = str(row["description"]).strip()
        if description in allowed and description not in catalog:
            catalog[description] = {"position_id": row["position_id"], "description": description}
    return catalog


def _requested_approval_descriptions(position_key: str, ai_level: int) -> list[str]:
    configured = APPROVAL_POSITION_DESCRIPTIONS.get(position_key)
    if configured is None:
        return []

    if isinstance(configured, str):
        return [configured]

    capped_level = max(0, min(ai_level, MAX_AUTO_APPROVAL_LEVEL, max(configured.keys())))
    return [configured[level] for level in sorted(configured) if level <= capped_level]


def _extract_recommended_positions(ai_payload: dict[str, Any], catalog: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    recommendations: list[dict[str, Any]] = []
    ai_positions = ai_payload.get("positions", {}) if isinstance(ai_payload, dict) else {}

    for key, details in ai_positions.items():
        if key not in APPROVAL_POSITION_DESCRIPTIONS:
            continue
        if not isinstance(details, dict):
            continue

        ai_status = str(details.get("status", "no_experience"))
        ai_level = LEVEL_MAP.get(ai_status, 0)
        if ai_level < 1:
            continue

        base_name = POSITION_KEY_TO_NAME.get(key, str(key).replace("_", " ").title())
        approval_level = min(ai_level, MAX_AUTO_APPROVAL_LEVEL)
        qualified_label = f"{base_name} {approval_level}" if key in LEVELED_POSITION_KEYS else base_name
        for description in _requested_approval_descriptions(key, ai_level):
            matched = catalog.get(description)
            recommendations.append(
                {
                    "key": key,
                    "label": qualified_label,
                    "base_name": base_name,
                    "ai_level": ai_level,
                    "ai_status": ai_status,
                    "estimated_years": details.get("estimated_years"),
                    "confidence": details.get("confidence"),
                    "reasons": details.get("reasons") if isinstance(details.get("reasons"), list) else [],
                    "position_id": matched.get("position_id") if matched else None,
                    "position_description": matched.get("description") if matched else description,
                    "match_found": matched is not None,
                }
            )

    recommendations.sort(key=lambda item: (item.get("base_name", ""), item.get("position_description", "")))
    return recommendations


def _get_existing_eligible_position_ids(conn, employee_id: int) -> set[int]:
    rows = conn.execute(
        text(
            """
            SELECT position_id
            FROM employee_position
            WHERE employee_id = :employee_id
              AND eligible = 1
              AND (status = 1 OR status = 'ACTIVE')
            """
        ),
        {"employee_id": employee_id},
    ).fetchall()
    return {int(row[0]) for row in rows if row[0] is not None}


def _get_new_active_employees(target_date: str) -> list[dict[str, Any]]:
    engine = _engine()
    try:
        with engine.connect() as conn:
            rows = conn.execute(
                text(
                    """
                    SELECT employee_id, first_name, last_name, email, mobile, home, work,
                           resume, start_date, start_date2
                    FROM employee
                    WHERE status = 1
                      AND deleted_at IS NULL
                      AND resume IS NOT NULL
                      AND resume != ''
                      AND (DATE(start_date) = :today OR DATE(start_date2) = :today)
                    ORDER BY employee_id
                    """
                ),
                {"today": target_date},
            ).mappings().all()
            return [dict(row) for row in rows]
    finally:
        engine.dispose()


async def _extract_resume_text_or_image_message(filename: str) -> tuple[str | None, list[dict[str, Any]] | None]:
    resume_url = _resume_url(filename)
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(resume_url, follow_redirects=True)
    if resp.status_code != 200:
        raise RuntimeError(f"Could not download resume from S3. Status {resp.status_code}.")

    content = resp.content
    filename_lower = filename.lower()
    mime = mimetypes.guess_type(filename)[0] or ""

    if filename_lower.endswith(".pdf") or "pdf" in mime:
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted:
                pages.append(extracted)
        return "\n\n".join(pages).strip(), None

    if filename_lower.endswith(".docx") or "wordprocessingml.document" in mime:
        document = docx.Document(io.BytesIO(content))
        parts = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)
        return "\n\n".join(parts).strip(), None

    if mime.startswith("image/") or filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        encoded = base64.b64encode(content).decode("utf-8")
        safe_mime = mime if mime.startswith("image/") else "image/png"
        messages = [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Transcribe this resume image and evaluate it using the provided schema. Respond with JSON only.",
                    },
                    {"type": "image_url", "image_url": {"url": f"data:{safe_mime};base64,{encoded}", "detail": "low"}},
                ],
            },
        ]
        return None, messages

    try:
        return content.decode("utf-8").strip(), None
    except UnicodeDecodeError as exc:
        raise RuntimeError("Unsupported resume file format for automatic text extraction.") from exc


async def _analyze_resume(filename: str) -> tuple[dict[str, Any], str]:
    from apps.api_usage_tracker import log_api_usage
    log_api_usage("New Employee Position Approver")

    api_key = os.getenv("NEW_EMPLOYEE_POSITION_APPROVER") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("NEW_EMPLOYEE_POSITION_APPROVER or OPENAI_API_KEY is not configured.")

    resume_text, image_messages = await _extract_resume_text_or_image_message(filename)
    if image_messages:
        messages = image_messages
    else:
        if not resume_text:
            raise RuntimeError("No text could be extracted from the resume.")
        messages = [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"Resume extracted from {filename}. Return only the JSON schema provided.\n\n{resume_text[:20000]}",
            },
        ]

    client = openai.AsyncOpenAI(api_key=api_key)
    response = await client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL_AUTOMATION", "gpt-4.1-nano"),
        messages=messages,
        response_format={"type": "json_object"},
        temperature=0.2,
    )
    content = response.choices[0].message.content if response.choices else ""
    if not content:
        raise RuntimeError("AI response was empty.")

    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as exc:
        raise RuntimeError("AI response was not valid JSON.") from exc

    return parsed, resume_text or "Resume was evaluated from an image."


def _build_review_record(
    employee: dict[str, Any],
    target_date: str,
    ai_payload: dict[str, Any] | None,
    resume_text: str,
    recommendations: list[dict[str, Any]],
    status: str,
    error: str = "",
) -> dict[str, Any]:
    employee_id = int(employee["employee_id"])
    first_name = employee.get("first_name") or ""
    last_name = employee.get("last_name") or ""
    phone = employee.get("mobile") or employee.get("home") or employee.get("work") or ""
    resume_file = employee.get("resume") or ""
    summary = ai_payload.get("candidate_summary", {}) if isinstance(ai_payload, dict) else {}

    return {
        "review_id": f"{target_date}:{employee_id}",
        "review_date": target_date,
        "employee_id": employee_id,
        "employee_name": f"{first_name} {last_name}".strip() or f"Employee {employee_id}",
        "first_name": first_name,
        "last_name": last_name,
        "email": employee.get("email") or "",
        "phone": phone,
        "hire_date": str(employee.get("start_date") or ""),
        "rehire_date": str(employee.get("start_date2") or ""),
        "resume_file": resume_file,
        "resume_url": _resume_url(resume_file) if resume_file else "",
        "status": status,
        "recommendations": recommendations,
        "ai_summary": summary,
        "ai_payload": ai_payload or {},
        "resume_text_preview": resume_text[:5000],
        "error": error,
        "created_at": _now_iso(),
        "completed": False,
        "approved_positions": [],
        "dismissed_reason": "",
    }


async def scan_for_new_employees() -> dict[str, Any]:
    target_date = _today()

    async with _state_lock:
        state = load_state()
        reviewed_today = set(map(str, state["reviewed"].get(target_date, [])))

    employees = _get_new_active_employees(target_date)
    candidates = [emp for emp in employees if str(emp["employee_id"]) not in reviewed_today]
    if not candidates:
        async with _state_lock:
            state = load_state()
            state["last_scan_at"] = _now_iso()
            state["last_scan_message"] = f"No new active employees with resumes found for {target_date}."
            save_state(state)
        return {"checked": 0, "added": 0, "message": state["last_scan_message"]}

    engine = _engine()
    added = 0
    checked = 0
    records_to_add: list[dict[str, Any]] = []
    reviewed_ids: list[str] = []

    try:
        with engine.connect() as conn:
            catalog = _get_allowed_catalog(conn)

            for employee in candidates:
                checked += 1
                employee_id = int(employee["employee_id"])
                reviewed_ids.append(str(employee_id))
                try:
                    ai_payload, resume_text = await _analyze_resume(employee.get("resume") or "")
                    recommendations = _extract_recommended_positions(ai_payload, catalog)
                    existing_ids = _get_existing_eligible_position_ids(conn, employee_id)
                    recommendations = [
                        item for item in recommendations if item.get("position_id") not in existing_ids
                    ]
                    
                    status = "Recommended" if recommendations else "No New Positions"
                    if recommendations:
                        added += 1
                        
                    # --- Auto-Approve Intercept ---
                    auto_approved_pos = []
                    completed = False
                    if recommendations and get_automation_enabled():
                        match_recs = [r for r in recommendations if r.get("position_id")]
                        if match_recs:
                            try:
                                auto_approved_pos = _execute_approval_db_logic(employee_id, match_recs)
                                if auto_approved_pos:
                                    status = "Approved"
                                    completed = True
                                    # Dispatch email
                                    email_addr = employee.get("email")
                                    first_name = employee.get("first_name") or "Employee"
                                    if email_addr:
                                        send_position_added_email(email_addr, first_name, auto_approved_pos)
                            except Exception:
                                logger.exception(f"Failed to auto-approve positions for employee {employee_id}")

                    record = _build_review_record(
                        employee,
                        target_date,
                        ai_payload,
                        resume_text,
                        recommendations,
                        status,
                    )
                    if completed:
                        record["completed"] = True
                        record["approved_positions"] = auto_approved_pos
                        record["approved_at"] = _now_iso()
                except Exception as exc:
                    logger.exception("Failed reviewing employee %s", employee_id)
                    record = _build_review_record(
                        employee,
                        target_date,
                        None,
                        "",
                        [],
                        "Review Error",
                        str(exc),
                    )

                records_to_add.append(record)
                await asyncio.sleep(1)
    finally:
        engine.dispose()

    async with _state_lock:
        state = load_state()
        existing_review_ids = {record.get("review_id") for record in state.get("records", [])}
        for record in records_to_add:
            if record["review_id"] not in existing_review_ids:
                state["records"].insert(0, record)

        day_reviewed = set(map(str, state["reviewed"].get(target_date, [])))
        day_reviewed.update(reviewed_ids)
        state["reviewed"][target_date] = sorted(day_reviewed, key=lambda value: int(value) if value.isdigit() else value)
        state["last_scan_at"] = _now_iso()
        state["last_scan_message"] = f"Checked {checked} new employee(s); added {added} recommendation record(s)."
        save_state(state)

    return {"checked": checked, "added": added, "message": state["last_scan_message"]}


def send_position_added_email(employee_email: str, first_name: str, added_positions: list[str]) -> bool:
    sender_email = "golive@culinarystaffing.com"
    
    import requests
    tenant_id = os.getenv("O365_TENANT_ID")
    client_id = os.getenv("O365_CLIENT_ID")
    client_secret = os.getenv("O365_CLIENT_SECRET")
    
    if not all([tenant_id, client_id, client_secret, employee_email]):
        logger.warning("Skipping position update email: Microsoft 365 OAuth credentials or email missing.")
        return False
        
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        "grant_type": "client_credentials",
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://graph.microsoft.com/.default"
    }
    
    try:
        r = requests.post(token_url, data=token_data, timeout=15)
        r.raise_for_status()
        access_token = r.json().get("access_token")
    except Exception as e:
        logger.exception("Failed to authenticate with Microsoft Graph for positions email")
        return False

    subject = "New Position(s) Added to Your GoLive! Profile"
    
    html_body = f"""
    <html>
      <body style="font-family: Arial, sans-serif; color: #333; max-width: 600px; margin: 0 auto; padding: 20px;">
        <h2 style="color: #047857;">Profile Update</h2>
        <p>Hi {first_name},</p>
        <p>Good news! Your request has been approved, and the following position(s) have been added to your profile:</p>
        <ul style="background: #f0fdf4; padding: 15px 35px; border-radius: 6px; border: 1px solid #dcfce7; list-style: square;">
    """
    for p in added_positions:
        html_body += f"<li style='margin-bottom: 4px;'><strong>{p}</strong></li>"
        
    html_body += f"""
        </ul>
        <p>You will now be eligible to request shifts for these positions in GoLive!</p>
        <p style="margin-top: 30px;">Best regards,<br>The Culinary Staffing Team</p>
      </body>
    </html>
    """
    
    email_msg = {
        "message": {
            "subject": subject,
            "body": {
                "contentType": "HTML",
                "content": html_body
            },
            "toRecipients": [
                {"emailAddress": {"address": employee_email}}
            ]
        },
        "saveToSentItems": "true"
    }

    send_url = f"https://graph.microsoft.com/v1.0/users/{sender_email}/sendMail"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    try:
        send_res = requests.post(send_url, headers=headers, json=email_msg, timeout=30)
        send_res.raise_for_status()
        return True
    except Exception as e:
        logger.exception("Failed to dispatch position email via MS Graph")
        return False

def _execute_approval_db_logic(employee_id: int, approved_recommendations: list[dict[str, Any]]) -> list[str]:
    """Internal helper to run direct DB SQL updates for matched positions."""
    engine = _engine()
    added_positions: list[str] = []
    try:
        with engine.begin() as conn:
            allowed_catalog = _get_allowed_catalog(conn)
            allowed_position_ids = {
                int(item["position_id"]) for item in allowed_catalog.values() if item.get("position_id")
            }
            
            valid_recs = [
                item for item in approved_recommendations 
                if item.get("position_id") and int(item["position_id"]) in allowed_position_ids
            ]
            
            for item in valid_recs:
                position_id = int(item["position_id"])
                existing = conn.execute(
                    text("""
                        SELECT employee_position_id
                        FROM employee_position
                        WHERE employee_id = :employee_id AND position_id = :position_id
                        LIMIT 1
                    """),
                    {"employee_id": employee_id, "position_id": position_id},
                ).fetchone()

                if existing:
                    conn.execute(
                        text("""
                            UPDATE employee_position
                            SET eligible = 1, status = 1
                            WHERE employee_position_id = :employee_position_id
                        """),
                        {"employee_position_id": existing[0]},
                    )
                else:
                    conn.execute(
                        text("""
                            INSERT INTO employee_position (employee_id, position_id, eligible, status, sub_type_id)
                            VALUES (:employee_id, :position_id, 1, 1, -1)
                        """),
                        {"employee_id": employee_id, "position_id": position_id},
                    )

                added_positions.append(item.get("position_description") or item.get("label") or str(position_id))
    finally:
        engine.dispose()
        
    return added_positions


def approve_positions_for_review(review_id: str, selected_position_ids: list[int]) -> dict[str, Any]:
    selected_ids = {int(item) for item in selected_position_ids if item}
    if not selected_ids:
        return {"status": "error", "message": "Select at least one matched position to approve."}

    state = load_state()
    record = next((item for item in state.get("records", []) if item.get("review_id") == review_id), None)
    if not record:
        return {"status": "error", "message": "Review record not found."}

    recommendations = record.get("recommendations", [])
    approved_recommendations = [
        item for item in recommendations if item.get("position_id") and int(item["position_id"]) in selected_ids
    ]
    if not approved_recommendations:
        return {"status": "error", "message": "No matched recommended positions were selected."}

    employee_id = int(record["employee_id"])
    try:
        added_positions = _execute_approval_db_logic(employee_id, approved_recommendations)
        if not added_positions:
            return {
                "status": "error",
                "message": "The selected positions are no longer in the approved position catalog.",
            }
            
        # Attempt Email Dispatch
        employee_email = record.get("email")
        first_name = record.get("first_name") or "Employee"
        if employee_email:
            send_position_added_email(employee_email, first_name, added_positions)
            
    except Exception as e:
        logger.exception(f"Database operation failed approving positions for {employee_id}")
        return {"status": "error", "message": f"Internal execution error: {e}"}

    record["status"] = "Approved"
    record["completed"] = True
    record["approved_positions"] = added_positions
    record["approved_at"] = _now_iso()
    save_state(state)
    return {
        "status": "success",
        "message": f"Approved {len(added_positions)} position(s): {', '.join(added_positions)}",
    }


def refresh_pending_records_against_allowed_catalog() -> None:
    """Update older pending records so they use only the current approval catalog."""
    state = load_state()
    records = state.get("records", [])
    if not records:
        return

    changed = False
    engine = _engine()
    try:
        with engine.connect() as conn:
            catalog = _get_allowed_catalog(conn)
            for record in records:
                if record.get("completed") or not record.get("ai_payload"):
                    continue

                employee_id = int(record["employee_id"])
                recommendations = _extract_recommended_positions(record["ai_payload"], catalog)
                existing_ids = _get_existing_eligible_position_ids(conn, employee_id)
                recommendations = [
                    item for item in recommendations if item.get("position_id") not in existing_ids
                ]

                if recommendations != record.get("recommendations", []):
                    record["recommendations"] = recommendations
                    record["status"] = "Recommended" if recommendations else "No New Positions"
                    record["catalog_refreshed_at"] = _now_iso()
                    changed = True
    except Exception:
        logger.exception("Could not refresh pending New Employee Position Approver records")
        return
    finally:
        engine.dispose()

    if changed:
        save_state(state)


def dismiss_review(review_id: str, reason: str = "") -> dict[str, Any]:
    state = load_state()
    record = next((item for item in state.get("records", []) if item.get("review_id") == review_id), None)
    if not record:
        return {"status": "error", "message": "Review record not found."}

    record["status"] = "Dismissed"
    record["completed"] = True
    record["dismissed_reason"] = reason
    record["dismissed_at"] = _now_iso()
    save_state(state)
    return {"status": "success", "message": "Review dismissed."}


async def new_employee_position_approver_loop() -> None:
    while True:
        try:
            await scan_for_new_employees()
        except asyncio.CancelledError:
            break
        except Exception as exc:
            logger.exception("New Employee Position Approver loop failed")
            async with _state_lock:
                state = load_state()
                state["last_scan_at"] = _now_iso()
                state["last_scan_message"] = f"Scan error: {exc}"
                save_state(state)

        await asyncio.sleep(60 * 15)
