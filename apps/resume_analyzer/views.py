from __future__ import annotations

import base64
import io
import json
import logging
import os
from typing import List

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from openai import OpenAI, OpenAIError
from pypdf import PdfReader
from docx import Document

RESUME_SYSTEM_PROMPT = """
You are a resume screener for a hospitality staffing agency.
The user will send you a resume (as text or transcribed from a PDF/Word doc/image).
Your job is to decide how qualified the candidate is for specific hospitality positions. 
Count experience at fine-dining or equivalent venues normally, but treat fast-food experience differently (see updated rules below).

Target positions

Evaluate the candidate for these positions:

Cook
Prep Cook
Dishwasher
Utility
Server
Host
Runner
Busser
Bartender
Barback
Cashier
Pastry
Baker
Sushi
Concessions
Barista
Valet
Event Supervisor

Venue rules (VERY IMPORTANT)

Fast-food experience no longer disqualifies a candidate. Instead:

— Fast food or clearly quick-service chains (e.g., McDonald’s, Burger King, Wendy’s, Taco Bell, KFC, In-N-Out, Chick-fil-A, similar chains) DO qualify,
  BUT ONLY for Level 1 and ONLY if the role performed directly matches one of the target positions.

Examples:
• A McDonald’s Cashier → qualifies for the Cashier position at Level 1.
• A Wendy’s Cook → qualifies for the Cook position at Level 1.
• A Taco Bell Crew Member with cashier duties → qualifies for Cashier Level 1.
• A fast-food Shift Lead → does NOT qualify unless duties explicitly match one of the listed roles.

Fast-food experience should never count toward Level 2 or Level 3.

For Level 2 or Level 3 qualification:
Count only experience at fine dining or equivalent hospitality venues, such as:
— Hotels, resorts, country clubs
— Upscale restaurants, steakhouses, chef-driven or white-tablecloth concepts
— Banquet/catering companies, convention centers, stadiums, arenas, large event venues
— Corporate/contract dining for companies, universities, hospitals, etc., when clearly hospitality-related.

If a venue type is unclear and could reasonably be hospitality (e.g., “Italian restaurant” without branding), you may count it with reduced confidence.

Ignore non-hospitality jobs entirely (admin, warehouse, rideshare, retail, etc.).

Special rule for Event Supervisor:
— This position requires a minimum of 3 years of management or supervisory experience in the hotel, food & beverage, or hospital industry.

Experience rules

For each position, you must:
1. Examine the entire work history and identify matching roles.
2. Estimate total time (in years) spent in those roles.

Experience categorization:
• Level 1: less than 2 years combined qualifying experience.
  — All fast-food experience ALWAYS counts as Level 1.
• Level 2: 2 to 5 years combined qualifying experience at non-fast-food venues.
• Level 3: more than 5 years qualifying experience at non-fast-food venues.

If the candidate has only fast-food experience for a role, assign Level 1 (never “no_experience”).

If they have neither qualifying nor fast-food experience, assign "no_experience".

When estimating experience:
— Use job dates when available.
— Estimate approximate duration when missing.
— Avoid double counting overlapping jobs.

Output format

Return your result as valid JSON only, using this schema:
{
  "candidate_summary": {
    "hospitality_experience_overview": "",
    "total_hospitality_years_estimate": 0.0,
    "notable_venues": [],
    "notes_on_fast_food_or_non_qualifying_experience": ""
  },
  "positions": {
    "cook": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "prep_cook": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "dishwasher": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "utility": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "server": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "host": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "runner": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "busser": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "bartender": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "barback": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "cashier": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "pastry": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "baker": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "sushi": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "concessions": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "barista": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "valet": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    },
    "event_supervisor": {
      "status": "no_experience | level_1 | level_2 | level_3",
      "estimated_years": 0.0,
      "confidence": 0.0,
      "reasons": []
    }
  }
}

Confidence must be between 0.0 and 1.0.

In the reasons field, briefly explain:
— which roles and venues you counted,
— if fast-food experience was used to assign Level 1,
— why any other roles were excluded.

Do not include any text outside of the JSON.
"""



logger = logging.getLogger(__name__)
templates = Jinja2Templates(directory="templates")
router = APIRouter()


@router.get("", response_class=HTMLResponse)
async def page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse("apps/resume_analyzer.html", {"request": request})


@router.post("/analyze", response_class=JSONResponse)
async def analyze_resume(
    file: UploadFile | None = File(None),
    resume_text: str | None = Form(None),
) -> JSONResponse:
    if file:
        contents = await file.read()
        if not contents:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    else:
        contents = b""

    try:
        client = _get_openai_client()
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    try:
        ai_payload = _build_resume_messages(
            file.filename if file else "",
            file.content_type if file else "",
            contents,
            resume_text,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=ai_payload,
            response_format={"type": "json_object"},
        )
    except OpenAIError as exc:  # pragma: no cover - external API call
        logger.exception("OpenAI API request failed")
        detail = getattr(exc, "message", None) or getattr(exc, "error", None) or str(exc)
        raise HTTPException(status_code=502, detail=f"AI service error: {detail}") from exc
    except Exception as exc:  # pragma: no cover - external API call
        logger.exception("Unexpected error while contacting AI service")
        raise HTTPException(status_code=502, detail="Unable to reach the AI service.") from exc

    content_block = response.choices[0].message.content if response.choices else ""
    if not content_block:
        raise HTTPException(status_code=502, detail="AI response was empty.")

    try:
        parsed = json.loads(content_block)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="AI response was not valid JSON.") from exc

    return JSONResponse(parsed)


def _get_openai_client() -> OpenAI:
    api_key = os.getenv("RESUME_ANALYZER_OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "RESUME_ANALYZER_OPENAI_API_KEY is not configured. Falling back to OPENAI_API_KEY is also not set."
        )
    return OpenAI(api_key=api_key)


def _build_resume_messages(
    filename: str,
    content_type: str,
    contents: bytes,
    resume_text: str | None,
) -> list[dict]:
    name = (filename or "uploaded file").strip()
    mime = (content_type or "").lower()
    lower_name = name.lower()
    text_input = (resume_text or "").strip()

    if "pdf" in mime or lower_name.endswith(".pdf"):
        text = _extract_pdf_text(contents)
        if not text:
            raise ValueError("Could not read any text from the uploaded PDF.")
        return [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"Resume extracted from {name}. Return only the JSON schema provided.\n\n{text}",
            },
        ]

    if _is_docx_file(mime, lower_name):
        text = _extract_docx_text(contents)
        if not text:
            raise ValueError("Could not read any text from the uploaded Word document.")
        return [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"Resume extracted from {name}. Return only the JSON schema provided.\n\n{text}",
            },
        ]

    if _is_image_file(mime, lower_name):
        data_url = _encode_image_to_data_url(contents, mime)
        return [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Transcribe this resume image and evaluate it using the provided schema. Respond with JSON only.",
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ]

    if text_input:
        return [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"Resume provided as pasted text. Return only the JSON schema provided.\n\n{text_input}",
            },
        ]

    raise ValueError("Upload a PDF, Word document, image, or paste resume text.")


def _extract_pdf_text(file_bytes: bytes) -> str:
    buffer = io.BytesIO(file_bytes)
    try:
        reader = PdfReader(buffer)
    except Exception as exc:  # pragma: no cover - external library
        raise ValueError(f"Could not read PDF: {exc}") from exc

    texts: List[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # pragma: no cover - library-specific errors
            text = ""
        if text:
            texts.append(text)
    return "\n\n".join(texts).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    buffer = io.BytesIO(file_bytes)
    try:
        document = Document(buffer)
    except Exception as exc:  # pragma: no cover - external library
        raise ValueError(f"Could not read Word document: {exc}") from exc

    texts = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    texts.append(cell_text)

    return "\n\n".join(texts).strip()


def _is_docx_file(mime: str, lower_name: str) -> bool:
    return (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower_name.endswith(".docx")
    )


def _is_image_file(mime: str, lower_name: str) -> bool:
    return mime.startswith("image/") or lower_name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp"))


def _encode_image_to_data_url(contents: bytes, mime: str) -> str:
    safe_mime = mime if mime.startswith("image/") else "image/png"
    encoded = base64.b64encode(contents).decode("utf-8")
    return f"data:{safe_mime};base64,{encoded}"
